import argparse
import psycopg2
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


# ---------------- DATABASE CONFIG ----------------
DB_CONFIG = {
    "dbname": "ShritanuDB",
    "user": "postgres",
    "password": "Smpm@2397",
    "host": "localhost",
    "port": "5432"
}

# --- HELPER: Determine previous years (Modified) ---
def get_previous_years(prediction_year, num_years=2):
    """
    Calculates the preceding financial years based on the prediction year.
    Handles 'YYYY-YY' and 'YYYY-YYYY' formats.
    
    Args:
        prediction_year (str): The financial year (e.g., "2024-25").
        num_years (int): How many previous years to calculate.
        
    Returns:
        list: A list of the preceding financial year strings.
    """
    previous_years = [] # 1. Create a list to store the results
    
    if len(prediction_year) == 7: # Handles 'YYYY-YY'
        try:
            base_year = int(prediction_year.split('-')[0])
            
            # 2. Loop 'num_years' times
            for i in range(1, num_years + 1):
                # i will be 1, then 2, then 3...
                prev_start = base_year - i
                prev_end_year = base_year - (i - 1)
                prev_end = str(prev_end_year)[-2:]
                
                previous_years.append(f"{prev_start}-{prev_end}")
                
        except (ValueError, IndexError):
            raise ValueError("Prediction year must be in 'YYYY-YY' format and parsable.")
            
    elif len(prediction_year) == 9: # Handles 'YYYY-YYYY'
        try:
            base_year = int(prediction_year.split('-')[0])
            
            # 2. Loop 'num_years' times
            for i in range(1, num_years + 1):
                prev_start = base_year - i
                prev_end = base_year - (i - 1)
                
                previous_years.append(f"{prev_start}-{prev_end}")
                
        except (ValueError, IndexError):
            raise ValueError("Prediction year must be in 'YYYY-YYYY' format and parsable.")
    else:
        raise ValueError("Prediction year must be in 'YYYY-YY' or 'YYYY-YYYY' format.")
    
    return previous_years # 3. Return the entire list

# ---------------- FETCH DATA (with Season) ----------------
def fetch_data(prediction_year, season):
    """
    Fetches data for a specific year AND season.
    The query already fetches all crops and states.
    """
    try:
        prev_year1, prev_year2 = get_previous_years(prediction_year)
    except ValueError as e:
        print(f"🔴 YEAR PARSING ERROR: {e}. Report generation aborted.")
        return {}, []

    target_years = [prediction_year, prev_year1, prev_year2]

    print(f"\n--- DB Query Debug ---")
    print(f"Target Years: {', '.join(target_years)} | Target Season: {season}")
    
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        
        # --- 1. Fetch ALL distinct prediction methods (for the specific year AND season) ---
        cur.execute("""
            SELECT DISTINCT method 
            FROM "NPCYF".crop_yields
            WHERE year = %s 
              AND method != 'MoA&FW' 
              AND season = %s
            ORDER BY method
        """, (prediction_year, season)) 
        
        prediction_methods = [row[0] for row in cur.fetchall()]
        
  
        cur.execute("""
            SELECT c.corp_name, s.state_name, cy.year, cy.method, cy.yield_value, cy.rmse_value
            FROM "NPCYF".crop_yields cy
            JOIN "NPCYF".crops c ON c.corp_id = cy.crop_id
            JOIN "NPCYF".states s ON s.state_id = cy.state_id
            WHERE cy.year = ANY(%s) AND cy.season = %s
            ORDER BY c.corp_name, s.state_name, cy.year DESC, cy.method
        """, (target_years, season)) 

        data = defaultdict(lambda: defaultdict(lambda: defaultdict(dict)))
        for crop, state, year, method, yield_val, rmse_val in cur.fetchall():
            data[crop.lower()][state][str(year)][method] = (yield_val, rmse_val)

        cur.close()
        conn.close()
        # print("--- End DB Query Debug ---\n")
        
        return data, prediction_methods

    except psycopg2.OperationalError as e:
        print(f"🔴 DATABASE CONNECTION ERROR: {e}. Report generation aborted.")
        return {}, []
    # except psycopg2.Error as e:
    #     print(f"🔴 DATABASE QUERY ERROR: {e}. Check if the 'season' column exists and matches.")
    #     return {}, []


# ---------------- DYNAMIC HEADER BUILDER ----------------
def build_dynamic_headers(prediction_year, dynamic_methods):
    """
    Builds the complete header structure dynamically.
    (This function does not need to change)
    """
    try:
        prev_year1, prev_year2 = get_previous_years(prediction_year)
    except ValueError:
        return []

    header_groups = []
    
    for method in dynamic_methods:
        header_groups.append((prediction_year, method, ['Yield', 'RMSE']))

    header_groups.append((prev_year1, 'MoA&FW', [])) 
    header_groups.append((prev_year2, 'MoA&FW', []))

    return header_groups


# ---------------- UTILITY FUNCTIONS ----------------
def set_cell_text(cell, text, bold=False, size=Pt(9)):
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(str(text) if text is not None else "") 
    run.bold = bold
    run.font.size = size
    
    try:
        tcPr = cell._element.get_or_add_tcPr() 
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
    except Exception as e:
        pass 

def draw_horizontal_rule(doc):
    hr = doc.add_paragraph()
    p = hr._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_footer(doc):
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"© {datetime.now().year} IDEAS-TIH. All rights reserved. | Date: {datetime.now().strftime('%d-%m-%Y')} | 1")
    run.font.size = Pt(8)


# ---------------- REPORT GENERATOR (with Page Breaks) ----------------
def create_report(template_path, output_path, orientation, logo_path, prediction_year, season):
    
    data, prediction_methods = fetch_data(prediction_year, season)
    
    if not data:
        print(f"🔴 Report generation failed: No data could be processed for {prediction_year} ({season}).")
        return

    doc = Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.5)
        if orientation.upper() == "LANDSCAPE":
            section.orientation = 1
            section.page_width, section.page_height = section.page_height, section.page_width

    # --- Header Logo + Title (Printed once at the top) ---
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_para.add_run()
    try:
        logo_run.add_picture(logo_path, width=Inches(1.0))
    except Exception as e:
        print(f"Warning: Could not load logo image at {logo_path}. Error: {e}")

    title = doc.add_paragraph("IDEAS - Institute of Data Engineering, Analytics and Science Foundation")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(11)

    subtitle = doc.add_paragraph("ISI Kolkata | https://www.ideas-tih.org | +91 6289351800")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(8)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hrun = heading.add_run(f"Crop wise yield forecasts : {season.title()} Season {prediction_year}")
    hrun.bold = True
    hrun.font.size = Pt(11)

    draw_horizontal_rule(doc)
    # --- End of Header ---

    # --- Build header structure ONCE ---
    header_groups = build_dynamic_headers(prediction_year, prediction_methods)
    
    if not header_groups:
        print(f"🔴 ERROR: Could not build any table headers. Aborting.")
        return

    num_cols = 1 + sum((len(subs) if subs else 1) for _, _, subs in header_groups)

    # --- Loop over every crop found in the data ---
    all_crops = sorted(data.keys()) 
    num_crops = len(all_crops) # Get the total number of crops

    # Use enumerate to get the index (i)
    for i, crop_name in enumerate(all_crops):
        states = data[crop_name]
        
        if not states:
            print(f"🟡 WARNING: Skipping crop '{crop_name}' (no state data found).")
            continue
            
        print(f"Processing report for crop: {crop_name.title()}...")

        # --- Add Crop Title ---
        crop_para = doc.add_paragraph()
        crop_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crop_run = crop_para.add_run(crop_name.upper())
        crop_run.bold = True
        crop_run.font.size = Pt(10)
        crop_run.font.color.rgb = RGBColor(0, 80, 160)

        # --- Create a new table for THIS crop ---
        table = doc.add_table(rows=2, cols=num_cols)
        table.style = "Table Grid"

        # --- Build Table Headers (Row 0 and Row 1) ---
        row0, row1 = table.rows[0].cells, table.rows[1].cells
        row0[0].merge(row1[0])
        set_cell_text(row0[0], "State", bold=True, size=Pt(9))

        col_index = 1
        for year, model, subs in header_groups:
            if len(subs) > 1:
                if col_index + len(subs) - 1 >= num_cols: continue
                row0[col_index].merge(row0[col_index + len(subs) - 1])
                set_cell_text(row0[col_index], f"{year} {model}", bold=True, size=Pt(9))
                for sub in subs:
                    set_cell_text(row1[col_index], sub, bold=True, size=Pt(9))
                    col_index += 1
            else:
                if col_index >= num_cols: continue
                row0[col_index].merge(row1[col_index])
                set_cell_text(row0[col_index], f"{year}\n{model}", bold=True, size=Pt(9)) 
                col_index += 1

        # --- Fill Table Data for THIS crop ---
        sorted_states = sorted(states.keys())
        
        for state in sorted_states:
            years = states[state] 
            row = table.add_row().cells
            set_cell_text(row[0], state, size=Pt(8))
            c = 1 
            for year, model, subs in header_groups:
                val = years.get(year, {}).get(model, (None, None))
                
                if len(subs) > 1:
                    if c + 1 <= num_cols: 
                        set_cell_text(row[c], val[0], size=Pt(8)); c += 1 
                        set_cell_text(row[c], val[1], size=Pt(8)); c += 1 
                else:
                    if c < num_cols: 
                        set_cell_text(row[c], val[0], size=Pt(8)); c += 1
        
        # --- NEW: ADD PAGE BREAK ---
        # Instead of doc.add_paragraph(), we add a page break.
        # We add a page break AFTER the crop, unless it's the LAST crop.
        if i < num_crops - 1:
            doc.add_page_break()

    # --- Add Footer and Save (after all crops are processed) ---
    add_footer(doc)
    doc.save(output_path)
    print(f"✅ Report generated successfully → {output_path}")

# ---------------- MAIN (with Season) ----------------
# 
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate crop yield report.")
    
    parser.add_argument("-t", "--template", required=True, help="Template input path (not used but required)")
    parser.add_argument("-o", "--output", required=True, help="Output .docx file name")
    parser.add_argument("-f", "--format", choices=["LANDSCAPE", "PORTRAIT"], default="PORTRAIT", help="Page orientation")
    parser.add_argument("-l", "--logo", required=True, help="Path to logo image")
    parser.add_argument("-y", "--year", required=True, help="Prediction year (e.g., '2025-2026' or '2025-26')")
    parser.add_argument("-s", "--season", required=True, help="Season for the report (e.g., 'Kharif', 'Rabi','Zaid'))")

    args = parser.parse_args()
    
    create_report(args.template, args.output, args.format, args.logo, args.year, args.season)